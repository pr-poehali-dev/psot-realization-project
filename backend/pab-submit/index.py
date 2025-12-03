import json
import os
from typing import Dict, Any
from datetime import datetime
import psycopg2
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
import boto3

def create_word_document(pab_data: Dict) -> BytesIO:
    '''Создание Word документа ПАБ'''
    doc = Document()
    
    # Заголовок
    title = doc.add_heading('Регистрация ПАБ', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Основная информация
    doc.add_paragraph(f"Название листа: {pab_data['doc_number']}")
    doc.add_paragraph(f"Номер документа: {pab_data['doc_number']}")
    doc.add_paragraph(f"Дата: {pab_data['doc_date']}")
    doc.add_paragraph(f"ФИО проверяющего: {pab_data['inspector_fio']}")
    doc.add_paragraph(f"Должность проверяющего: {pab_data['inspector_position']}")
    doc.add_paragraph(f"Участок: {pab_data.get('location', '')}")
    doc.add_paragraph(f"Проверяемый объект: {pab_data.get('checked_object', '')}")
    doc.add_paragraph(f"Подразделение: {pab_data.get('department', '')}")
    
    doc.add_paragraph()
    
    # Наблюдения
    for idx, obs in enumerate(pab_data['observations'], 1):
        doc.add_heading(f'Наблюдение №{idx}', level=2)
        doc.add_paragraph(f"Описание: {obs['description']}")
        doc.add_paragraph(f"Категория: {obs.get('category', '')}")
        doc.add_paragraph(f"Вид условий и действий: {obs.get('conditions_actions', '')}")
        doc.add_paragraph(f"Опасные факторы: {obs.get('hazard_factors', '')}")
        doc.add_paragraph(f"Мероприятия: {obs['measures']}")
        doc.add_paragraph(f"Ответственный: {obs.get('responsible_person', '')}")
        doc.add_paragraph(f"Срок: {obs.get('deadline', '')}")
        doc.add_paragraph()
    
    # Сохранение в BytesIO
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    return file_stream

def handler(event: Dict[str, Any], context: Any) -> Dict[str, Any]:
    '''
    Сохранение ПАБ, создание Word документа и отправка email
    '''
    
    if event.get('httpMethod') == 'OPTIONS':
        return {
            'statusCode': 200,
            'headers': {
                'Access-Control-Allow-Origin': '*',
                'Access-Control-Allow-Methods': 'POST, OPTIONS',
                'Access-Control-Allow-Headers': 'Content-Type, X-User-Id',
                'Access-Control-Max-Age': '86400'
            },
            'body': '',
            'isBase64Encoded': False
        }
    
    body = json.loads(event.get('body', '{}'))
    
    dsn = os.environ.get('DATABASE_URL')
    conn = psycopg2.connect(dsn)
    cur = conn.cursor()
    
    # Сохранение записи ПАБ
    cur.execute(
        """INSERT INTO pab_records 
        (doc_number, doc_date, inspector_fio, inspector_position, user_id, 
        department, location, checked_object, photo_url) 
        VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s) 
        RETURNING id""",
        (
            body['doc_number'],
            body['doc_date'],
            body['inspector_fio'],
            body['inspector_position'],
            body.get('user_id'),
            body.get('department'),
            body.get('location'),
            body.get('checked_object'),
            body.get('photo_url')
        )
    )
    
    pab_id = cur.fetchone()[0]
    
    # Сохранение наблюдений
    for obs in body['observations']:
        cur.execute(
            """INSERT INTO pab_observations 
            (pab_record_id, observation_number, description, category, 
            conditions_actions, hazard_factors, measures, responsible_person, deadline) 
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)""",
            (
                pab_id,
                obs['observation_number'],
                obs['description'],
                obs.get('category'),
                obs.get('conditions_actions'),
                obs.get('hazard_factors'),
                obs['measures'],
                obs.get('responsible_person'),
                obs.get('deadline')
            )
        )
    
    conn.commit()
    
    # Создание Word документа
    word_doc = create_word_document(body)
    doc_filename = f"{body['doc_number']}. {body['doc_date']}.docx"
    
    # Сохранение в S3 (имитация - в реальности нужен S3 клиент)
    # Здесь должна быть логика загрузки в storage_folders
    
    # Отправка email
    try:
        smtp_host = os.environ.get('SMTP_HOST')
        smtp_port = int(os.environ.get('SMTP_PORT', 587))
        smtp_user = os.environ.get('SMTP_USER')
        smtp_password = os.environ.get('SMTP_PASSWORD')
        admin_email = os.environ.get('ADMIN_EMAIL')
        
        msg = MIMEMultipart()
        msg['From'] = smtp_user
        msg['To'] = f"{body.get('responsible_email', '')}, {admin_email}"
        msg['Subject'] = f"Новая регистрация ПАБ: {body['doc_number']}"
        
        email_body = f"""
        Зарегистрирован новый ПАБ
        
        Номер: {body['doc_number']}
        Дата: {body['doc_date']}
        Проверяющий: {body['inspector_fio']}
        Подразделение: {body.get('department', '')}
        
        Во вложении находится полный документ.
        """
        
        msg.attach(MIMEText(email_body, 'plain', 'utf-8'))
        
        # Прикрепление Word документа
        part = MIMEBase('application', 'vnd.openxmlformats-officedocument.wordprocessingml.document')
        part.set_payload(word_doc.read())
        encoders.encode_base64(part)
        part.add_header('Content-Disposition', f'attachment; filename={doc_filename}')
        msg.attach(part)
        
        server = smtplib.SMTP(smtp_host, smtp_port)
        server.starttls()
        server.login(smtp_user, smtp_password)
        server.send_message(msg)
        server.quit()
        
    except Exception as e:
        print(f"Email error: {str(e)}")
    
    cur.close()
    conn.close()
    
    return {
        'statusCode': 200,
        'headers': {
            'Content-Type': 'application/json',
            'Access-Control-Allow-Origin': '*'
        },
        'body': json.dumps({
            'success': True,
            'pab_id': pab_id,
            'doc_number': body['doc_number']
        }, ensure_ascii=False),
        'isBase64Encoded': False
    }
